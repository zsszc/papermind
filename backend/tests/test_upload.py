"""上传接口测试：大小限制（413）、扩展名白名单（400）、正常小文件上传成功。

说明：
- 通过 monkeypatch 把 MAX_UPLOAD_SIZE 改小来测试 413，无需真的生成 50MB 文件；
- papers 导入触发的后台处理线程（embedding/ChromaDB/LLM）全部 mock 掉；
- thesis 上传的 DocxParser.parse 也 mock 掉，保证测试快速、离线、稳定。
"""

import io
import shutil
import uuid
import zipfile
from pathlib import Path

import pytest
from docx import Document

from app.models import Paper, ThesisFile
from app.routers import papers as papers_router
from app.routers import thesis as thesis_router
from app.services.upload_validation import (
    UploadValidationError,
    validate_docx,
    validate_pdf,
)

# 项目根目录（路由内用 Path(__file__).resolve().parents[3] 定位，上传目录必须在其下，
# 否则路由里的 relative_to 会抛 ValueError）
PROJECT_ROOT = Path(papers_router.__file__).resolve().parents[3]


def _make_temp_dir(name: str) -> Path:
    """在项目根下创建独立临时目录（隐藏命名，不污染真实数据目录）。"""
    path = PROJECT_ROOT / f".pytest_tmp_{name}_{uuid.uuid4().hex[:8]}"
    path.mkdir(parents=True, exist_ok=True)
    return path

# 最小 PDF 字节流：即使解析失败，导入接口也会容错（metadata 记 parse_error）继续
MINIMAL_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
    b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj\n"
    b"trailer<</Root 1 0 R>>\n%%EOF\n"
)

def _minimal_docx() -> bytes:
    """用 python-docx 生成真实最小 DOCX，避免伪 ZIP 绕过测试门禁。"""
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("测试毕业论文", level=0)
    document.add_paragraph("正文内容")
    document.save(buffer)
    return buffer.getvalue()


VALID_DOCX = _minimal_docx()


def _zip_bytes(entries, compression=zipfile.ZIP_DEFLATED) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=compression) as archive:
        for name, payload in entries:
            archive.writestr(name, payload)
    return buffer.getvalue()


@pytest.fixture()
def papers_env(client, monkeypatch):
    """papers 导入环境：数据目录重定向到临时目录，后台处理 mock 掉。"""
    papers_dir = _make_temp_dir("papers")
    notes_dir = _make_temp_dir("notes")
    monkeypatch.setattr(papers_router, "get_papers_dir", lambda: papers_dir)
    monkeypatch.setattr(papers_router, "get_notes_dir", lambda: notes_dir)
    # mock 后台处理入口，避免触发 embedding/ChromaDB/LLM
    monkeypatch.setattr(papers_router, "_process_paper_background", lambda paper_id: None)
    try:
        yield client, papers_dir, notes_dir
    finally:
        shutil.rmtree(papers_dir, ignore_errors=True)
        shutil.rmtree(notes_dir, ignore_errors=True)


@pytest.fixture()
def thesis_env(client, monkeypatch):
    """thesis 上传环境：数据目录重定向到临时目录，DocxParser.parse mock 掉。"""
    thesis_dir = _make_temp_dir("my-thesis")
    monkeypatch.setattr(thesis_router, "get_thesis_dir", lambda: thesis_dir)

    def fake_parse(self, path):
        return {
            "title": "测试毕业论文",
            "chapters": [],
            "word_count": 100,
            "citations": [],
            "paragraphs": [],
        }

    monkeypatch.setattr(thesis_router.DocxParser, "parse", fake_parse)
    try:
        yield client, thesis_dir
    finally:
        shutil.rmtree(thesis_dir, ignore_errors=True)


def test_import_oversized_pdf_returns_413(papers_env, monkeypatch):
    """超过单文件大小上限的 PDF 应返回 413，且不落地残留文件。"""
    client, papers_dir, _ = papers_env
    # 把上限改小到 10 字节，用小文件即可触发 413，无需生成 50MB
    monkeypatch.setattr(papers_router, "MAX_UPLOAD_SIZE", 10)

    r = client.post(
        "/api/papers/import",
        files=[("files", ("big.pdf", MINIMAL_PDF, "application/pdf"))],
    )
    assert r.status_code == 413
    assert list(papers_dir.iterdir()) == []


def test_import_invalid_extension_returns_400(papers_env):
    """非 .pdf 扩展名应返回 400。"""
    client, _, _ = papers_env
    r = client.post(
        "/api/papers/import",
        files=[("files", ("notes.txt", b"hello", "text/plain"))],
    )
    assert r.status_code == 400


def test_import_small_pdf_success(papers_env):
    """正常小 PDF 上传成功：落盘、建库记录、生成空笔记文件。"""
    client, papers_dir, notes_dir = papers_env
    r = client.post(
        "/api/papers/import",
        files=[("files", ("paper.pdf", MINIMAL_PDF, "application/pdf"))],
    )
    assert r.status_code == 200
    data = r.json()
    assert data["total"] == 1
    item = data["items"][0]
    assert item["filename"] == "paper.pdf"

    # PDF 与笔记文件均已写入（重定向后的 tmp 目录）
    assert (papers_dir / "paper.pdf").exists()
    assert (notes_dir / f"{item['id']}.md").exists()


def test_thesis_upload_oversized_returns_413(thesis_env, monkeypatch):
    """超过单文件大小上限的 docx 应返回 413，且不落地残留文件。"""
    client, thesis_dir = thesis_env
    monkeypatch.setattr(thesis_router, "MAX_UPLOAD_SIZE", 10)

    r = client.post(
        "/api/thesis/upload",
        files={"file": ("thesis.docx", VALID_DOCX, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 413
    assert list(thesis_dir.iterdir()) == []


def test_thesis_upload_invalid_extension_returns_400(thesis_env):
    """非 .docx 扩展名应返回 400。"""
    client, _ = thesis_env
    r = client.post(
        "/api/thesis/upload",
        files={"file": ("thesis.pdf", MINIMAL_PDF, "application/pdf")},
    )
    assert r.status_code == 400


def test_thesis_upload_small_docx_success(thesis_env):
    """正常小 docx 上传成功（DocxParser.parse 已 mock）。"""
    client, thesis_dir = thesis_env
    r = client.post(
        "/api/thesis/upload",
        files={"file": ("thesis.docx", VALID_DOCX, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["filename"] == "thesis.docx"
    assert data["title"] == "测试毕业论文"
    assert (thesis_dir / "thesis.docx").exists()


def test_validate_pdf_rejects_spoofed_extension(tmp_path):
    path = tmp_path / "spoofed.pdf"
    path.write_bytes(b"junk before %PDF-1.7")

    with pytest.raises(UploadValidationError, match="PDF"):
        validate_pdf(path)


def test_import_spoofed_pdf_returns_400_without_orphans(papers_env, db):
    client, papers_dir, notes_dir = papers_env

    response = client.post(
        "/api/papers/import",
        files=[("files", ("spoofed.pdf", b"not-a-pdf", "application/pdf"))],
    )

    assert response.status_code == 400
    assert list(papers_dir.iterdir()) == []
    assert list(notes_dir.iterdir()) == []
    assert db.query(Paper).count() == 0


@pytest.mark.parametrize(
    ("entries", "kwargs", "message"),
    [
        (["[Content_Types].xml"], {}, "word/document.xml"),
        (["word/document.xml"], {}, "Content_Types"),
        (["[Content_Types].xml", "word/document.xml", "../escape.xml"], {}, "路径"),
        (["[Content_Types].xml", "word/document.xml", "/absolute.xml"], {}, "路径"),
        (["[Content_Types].xml", "word/document.xml", "word\\escape.xml"], {}, "路径"),
        (["[Content_Types].xml", "word/document.xml", "word/document.xml"], {}, "重复"),
        (["[Content_Types].xml", "word/document.xml", "extra.xml"], {"max_members": 2}, "成员数"),
    ],
)
def test_validate_docx_rejects_invalid_members(tmp_path, entries, kwargs, message):
    path = tmp_path / "invalid.docx"
    path.write_bytes(_zip_bytes([(name, b"x") for name in entries]))

    with pytest.raises(UploadValidationError, match=message):
        validate_docx(path, **kwargs)


def test_validate_docx_rejects_single_and_total_uncompressed_limits(tmp_path):
    path = tmp_path / "large.docx"
    path.write_bytes(_zip_bytes([
        ("[Content_Types].xml", b"x" * 16),
        ("word/document.xml", b"y" * 16),
    ]))

    with pytest.raises(UploadValidationError, match="单个"):
        validate_docx(path, max_member_size=15, max_total_size=100)
    with pytest.raises(UploadValidationError, match="总量"):
        validate_docx(path, max_member_size=20, max_total_size=31)


def test_validate_docx_rejects_excessive_compression_ratio(tmp_path):
    path = tmp_path / "bomb.docx"
    path.write_bytes(_zip_bytes([
        ("[Content_Types].xml", b"x"),
        ("word/document.xml", b"A" * 20_000),
    ]))

    with pytest.raises(UploadValidationError, match="压缩比"):
        validate_docx(path, max_compression_ratio=10)


def test_validate_real_minimal_docx(tmp_path):
    path = tmp_path / "valid.docx"
    path.write_bytes(VALID_DOCX)

    validate_docx(path)


def test_thesis_invalid_docx_returns_400_without_orphans(thesis_env, db):
    client, thesis_dir = thesis_env

    response = client.post(
        "/api/thesis/upload",
        files={"file": ("spoofed.docx", b"not-a-zip", "application/octet-stream")},
    )

    assert response.status_code == 400
    assert list(thesis_dir.iterdir()) == []
    assert db.query(ThesisFile).count() == 0


def test_thesis_parse_failure_cleans_file_and_database(thesis_env, db, monkeypatch):
    client, thesis_dir = thesis_env

    def fail_parse(self, path):
        raise ValueError("私密解析异常")

    monkeypatch.setattr(thesis_router.DocxParser, "parse", fail_parse)
    response = client.post(
        "/api/thesis/upload",
        files={"file": ("broken.docx", VALID_DOCX, "application/octet-stream")},
    )

    assert response.status_code == 500
    assert "私密解析异常" not in response.text
    assert list(thesis_dir.iterdir()) == []
    assert db.query(ThesisFile).count() == 0


def test_thesis_commit_failure_cleans_file_and_database(thesis_env, db, monkeypatch):
    client, thesis_dir = thesis_env

    def fail_commit():
        raise RuntimeError("私密数据库异常")

    monkeypatch.setattr(db, "commit", fail_commit)
    response = client.post(
        "/api/thesis/upload",
        files={"file": ("commit.docx", VALID_DOCX, "application/octet-stream")},
    )

    assert response.status_code == 500
    assert "私密数据库异常" not in response.text
    assert list(thesis_dir.iterdir()) == []
    assert db.query(ThesisFile).count() == 0
